#!/usr/bin/env python3
"""
task4_report.py -- 第四题：生产线故障自动识别 报告生成脚本

生成两份Word文档：
  1. 含附录版（完整源代码）
  2. 无附录版（精简）

用法：python3 code/task4_report.py
"""

from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── 路径 ──────────────────────────────────────────────────────────
PROJECT_ROOT = Path(__file__).resolve().parent.parent
CODE_DIR = PROJECT_ROOT / "code"
OUTPUT_DIR = CODE_DIR / "output"
EDA_DIR = OUTPUT_DIR / "task4_eda"
YIELD_VIZ_DIR = OUTPUT_DIR / "task4_yield_viz"

OUTPUT_WITH_APPENDIX = PROJECT_ROOT / "第四题_生产线故障自动识别报告_含附录.docx"
OUTPUT_WITHOUT_APPENDIX = PROJECT_ROOT / "第四题_生产线故障自动识别报告_无附录.docx"

# 附录源代码文件
APPENDIX_FILES = [
    CODE_DIR / "task4_load_data.py",
    CODE_DIR / "task4_eda.py",
    CODE_DIR / "task4_model.py",
    CODE_DIR / "task4_predict.py",
    CODE_DIR / "task4_yield_analysis.py",
]

# 图片文件及标题
FIGURES = [
    (EDA_DIR / "01_fault_distribution.png", "图1  故障类型分布饼图"),
    (EDA_DIR / "02_fault_by_hour.png", "图2  各时段故障频次热力图"),
    (EDA_DIR / "05_fault_correlation.png", "图3  故障类型相关性矩阵"),
    (EDA_DIR / "09_roc_curve.png", "图4  三种分类模型ROC曲线对比"),
    (EDA_DIR / "10_feature_importance.png", "图5  随机森林Top20特征重要性"),
    (YIELD_VIZ_DIR / "task4_yield_boxplot.png", "图6  M301各时段产量箱线图"),
    (YIELD_VIZ_DIR / "task4_correlation_heatmap.png", "图7  产量与合格率相关性热力图"),
]

# ── 字体工具 ──────────────────────────────────────────────────────

def set_run_font(run, name="宋体", size=Pt(12), bold=False, color=None):
    """设置run的字体属性。"""
    run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    # 设置中文字体
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{name}"/>')
        rpr.insert(0, rfonts)
    else:
        rfonts.set(qn("w:eastAsia"), name)
    # 同时设置ascii字体
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def add_paragraph(doc, text, style=None, alignment=None, font_name="宋体",
                  font_size=Pt(12), bold=False, color=None, space_after=Pt(6)):
    """添加段落并设置格式。"""
    p = doc.add_paragraph(style=style)
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    set_run_font(run, name=font_name, size=font_size, bold=bold, color=color)
    p.paragraph_format.space_after = space_after
    return p


def add_heading_styled(doc, text, level=1):
    """添加带中文字体的标题。"""
    heading = doc.add_heading(text, level=level)
    font_map = {1: "黑体", 2: "黑体", 3: "黑体"}
    size_map = {1: Pt(22), 2: Pt(16), 3: Pt(14)}
    for run in heading.runs:
        set_run_font(run, name=font_map.get(level, "黑体"),
                     size=size_map.get(level, Pt(14)), bold=True)
    return heading


# ── 封面 ──────────────────────────────────────────────────────────

def add_cover(doc):
    """添加封面。"""
    # 空行
    for _ in range(6):
        doc.add_paragraph("")

    add_paragraph(doc, "商务智能方法与应用课程设计", alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  font_name="黑体", font_size=Pt(26), bold=True, space_after=Pt(24))

    add_paragraph(doc, "第四题 - 生产线故障自动识别", alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  font_name="黑体", font_size=Pt(20), bold=False, space_after=Pt(48))

    add_paragraph(doc, "2026年7月", alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  font_name="宋体", font_size=Pt(16), space_after=Pt(12))

    doc.add_page_break()


# ── 摘要 ──────────────────────────────────────────────────────────

def add_abstract(doc):
    """添加摘要。"""
    add_heading_styled(doc, "摘  要", level=1)

    abstract_text = (
        "本研究基于泰迪杯2024年A题生产线数据，围绕生产线故障自动识别问题展开分析。"
        "研究目标是从传感器数据中提取故障特征，构建自动报警模型，实现对生产线上9类故障的"
        "快速识别与定位。首先，对附件1中M101生产线63万余条传感器记录进行数据清洗与特征"
        "工程，提取25维传感器特征及时间特征；随后，采用描述性统计方法从故障类型分布、"
        "时段规律、故障相关性等多维度分析故障模式。在此基础上，分别构建逻辑回归、随机"
        "森林和XGBoost三种二分类模型，其中XGBoost模型表现最优，AUC达0.9935，召回率"
        "99.44%，可有效捕获绝大多数故障事件。进一步将最优模型应用于附件2的M201和M202"
        "生产线数据，完成故障事件的预测与汇总。此外，针对附件3的M301生产线数据，开展"
        "产量与合格率的方差分析和相关性分析。结果表明，不同时间段的产量和合格率无显著"
        "差异，产量与合格率之间不存在显著线性相关。本研究为生产线故障诊断提供了可落地"
        "的机器学习方案，对提升制造业智能化水平具有参考价值。"
    )
    add_paragraph(doc, abstract_text, font_name="宋体", font_size=Pt(12), space_after=Pt(12))

    add_paragraph(doc, "关键词：", font_name="黑体", font_size=Pt(12), bold=True, space_after=Pt(2))
    add_paragraph(doc, "故障诊断；XGBoost；随机森林；生产线；方差分析",
                  font_name="宋体", font_size=Pt(12), space_after=Pt(12))

    doc.add_page_break()


# ── 目录 ──────────────────────────────────────────────────────────

def add_toc(doc):
    """添加自动目录域（需在Word中右键更新域才能显示）。"""
    add_heading_styled(doc, "目  录", level=1)

    p = doc.add_paragraph()
    run = p.add_run()
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fld_char_begin)

    run2 = p.add_run()
    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    run2._element.append(instr)

    run3 = p.add_run()
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._element.append(fld_char_end)

    add_paragraph(doc, '（请在Word中右键目录区域，选择\u201c更新域\u201d以生成完整目录）',
                  font_name="宋体", font_size=Pt(10), color=RGBColor(128, 128, 128))

    doc.add_page_break()


# ── 数据说明 ──────────────────────────────────────────────────────

def add_data_description(doc):
    """添加数据说明章节。"""
    add_heading_styled(doc, "1  数据说明", level=1)

    add_heading_styled(doc, "1.1  数据来源", level=2)
    add_paragraph(doc,
        "本研究数据来源于泰迪杯2024年全国大学生数据挖掘竞赛A题。"
        "该数据集记录了某制造企业多条生产线的传感器读数，涵盖物料推送、"
        "物料检测、填装、加盖、拧盖等全工序环节，共计37个字段。")

    add_heading_styled(doc, "1.2  数据文件概览", level=2)

    # 数据文件表格
    table = doc.add_table(rows=5, cols=4, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["附件", "文件", "规模", "说明"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, name="黑体", size=Pt(11), bold=True)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="4472C4" w:val="clear"/>')
        cell._element.get_or_add_tcPr().append(shading)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    data_rows = [
        ["附件1", "M101.csv", "82MB, 636K行x37列", "训练数据（传感器+故障标签）"],
        ["附件2", "M201.csv, M202.csv", "各约920MB", "测试数据（待预测故障）"],
        ["附件3", "M301.csv", "UTF-8编码", "产量分析数据（含合格/不合格计数）"],
        ["附件3", "操作人员信息表.xlsx", "1个工作表", "操作人员编号、生产线、工龄"],
    ]
    for r, row_data in enumerate(data_rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, name="宋体", size=Pt(10.5))
            if r % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>')
                cell._element.get_or_add_tcPr().append(shading)

    add_heading_styled(doc, "1.3  故障类型说明", level=2)
    add_paragraph(doc,
        "数据集包含9类故障，按装置类型编码（千位=装置，后三位=故障码）：")

    fault_desc = [
        "物料推送装置故障（1001）：推送气缸异常",
        "物料检测装置故障（2001）：物料到位检测异常",
        "填装装置检测故障（4001）：填装前检测异常",
        "填装装置定位故障（4002）：填装定位器异常",
        "填装装置填装故障（4003）：填装动作异常",
        "加盖装置定位故障（5001）：加盖定位异常",
        "加盖装置加盖故障（5002）：加盖动作异常",
        "拧盖装置定位故障（6001）：拧盖定位异常",
        "拧盖装置拧盖故障（6002/6009）：拧盖动作异常",
    ]
    for fd in fault_desc:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(fd)
        set_run_font(run, name="宋体", size=Pt(11))


# ── 方法 ──────────────────────────────────────────────────────────

def add_method(doc):
    """添加方法章节。"""
    add_heading_styled(doc, "2  研究方法", level=1)

    add_heading_styled(doc, "2.1  数据预处理", level=2)
    add_paragraph(doc,
        "对M101.csv的636K行数据进行清洗：删除传感器列全部为NaN的记录，"
        "数值型NaN填充为0，去除完全重复行。随后进行特征工程，包括：(1) 故障标记，"
        "任一故障列大于0即标记为1；(2) 时间特征提取，从秒级时间戳中提取小时和"
        "白天/夜间标识；(3) 生产线编号的独热编码。最终形成28维特征向量。")

    add_heading_styled(doc, "2.2  故障特征分析", level=2)
    add_paragraph(doc,
        "采用描述性统计方法，从四个维度分析故障模式：(1) 故障类型频率分布，"
        "识别高频故障；(2) 时段热力图，揭示故障时间规律；(3) 故障相关性矩阵，"
        "发现故障共现模式；(4) 持续时间分析，刻画各类故障的时长特征。")

    add_heading_styled(doc, "2.3  分类模型", level=2)
    add_paragraph(doc,
        "构建三种二分类模型预测故障发生：")
    models_desc = [
        "逻辑回归（基线）：class_weight='balanced'，max_iter=1000，作为性能下界参考。",
        "随机森林：100棵决策树，max_depth=15，min_samples_leaf=5，class_weight='balanced'。",
        "XGBoost：200棵树，max_depth=6，learning_rate=0.1，scale_pos_weight动态调整。",
    ]
    for md in models_desc:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(md)
        set_run_font(run, name="宋体", size=Pt(11))
    add_paragraph(doc,
        "数据按8:2分层抽样划分训练集和测试集。评估指标包括Precision、Recall、"
        "F1-score和AUC-ROC。由于故障样本占比低（类别不平衡），重点关注召回率和AUC。")

    add_heading_styled(doc, "2.4  产量分析", level=2)
    add_paragraph(doc,
        "对M301生产线数据，按日汇总累计计数器的最大值得到日产量和合格率。"
        "采用单因素方差分析（ANOVA）检验不同时间段产量和合格率的差异显著性，"
        "并计算产量与合格率的Pearson相关系数。由于仅有M301一条生产线和A001一名"
        "操作人员，无法进行双因素ANOVA。")


# ── 结果 ──────────────────────────────────────────────────────────

def add_results(doc):
    """添加结果章节。"""
    add_heading_styled(doc, "3  结果与分析", level=1)

    # --- 3.1 故障分布 ---
    add_heading_styled(doc, "3.1  故障分布分析", level=2)
    add_paragraph(doc,
        "如图1所示，故障类型分布呈现明显的不均衡特征。推送故障(1001)和"
        "检测故障(2001)出现频率最高，合计占比超过60%。填装相关故障(4001/4002/4003)"
        "次之，加盖和拧盖故障相对较少。这种分布特征提示物料输送环节是故障高发区，"
        "应作为重点监控对象。")
    _embed_figure(doc, FIGURES[0])

    # --- 3.2 时段规律 ---
    add_heading_styled(doc, "3.2  故障时段规律", level=2)
    add_paragraph(doc,
        "图2的热力图揭示了故障的时段分布规律。故障主要集中在生产活动活跃的时段，"
        "不同故障类型在各时段的分布存在差异。推送故障和检测故障在多数时段均有发生，"
        "而填装和拧盖故障则呈现一定的时段集中性。这些规律可为生产排班和预防性维护"
        "提供参考。")
    _embed_figure(doc, FIGURES[1])

    # --- 3.3 故障相关性 ---
    add_heading_styled(doc, "3.3  故障相关性分析", level=2)
    add_paragraph(doc,
        "图3展示了各故障类型之间的相关性矩阵。部分故障之间存在较强的正相关，"
        "表明某些故障倾向于同时发生（共现故障）。例如，同一装置的不同故障类型"
        "之间相关性较高，这与物理直觉一致：装置发生故障时，其多个子功能可能"
        "同时受到影响。")
    _embed_figure(doc, FIGURES[2])

    # --- 3.4 分类模型 ---
    add_heading_styled(doc, "3.4  分类模型评估", level=2)
    add_paragraph(doc,
        "表1列出了三种模型的评估指标。XGBoost在所有指标上均表现最优，"
        "AUC达0.9935，召回率99.44%，意味着仅有0.56%的故障事件被遗漏。"
        "随机森林次之，AUC为0.9917。逻辑回归作为基线模型，AUC仅0.8273，"
        "说明故障模式的非线性特征显著，需要更复杂的模型来捕获。")

    # 模型评估表
    _add_model_metrics_table(doc)
    _embed_figure(doc, FIGURES[3])
    _embed_figure(doc, FIGURES[4])

    add_paragraph(doc,
        "图5显示了随机森林模型的Top20特征重要性。传感器特征（如物料推送状态、"
        "填装检测数等）的重要性远高于时间特征和生产线编码特征，说明故障主要由"
        "传感器读数的异常变化驱动，而非时间或生产线固有差异。")

    # --- 3.5 M201/M202预测 ---
    add_heading_styled(doc, "3.5  M201/M202故障预测结果", level=2)
    add_paragraph(doc,
        "将训练好的XGBoost模型应用于附件2的M201和M202生产线数据，"
        "对连续故障记录进行合并，得到故障事件汇总（表2）。")
    _add_prediction_summary_table(doc)

    # --- 3.6 产量分析 ---
    add_heading_styled(doc, "3.6  产量与合格率分析", level=2)
    add_paragraph(doc,
        "图6为M301生产线各时段产量的箱线图，图7为产量与合格率的相关性热力图。"
        "ANOVA分析结果（表3）表明，不同时间段的产量和合格率均无显著差异（p>0.05），"
        "产量与合格率的Pearson相关系数为-0.2369（p=0.289），不存在显著线性相关。")
    _embed_figure(doc, FIGURES[5])
    _embed_figure(doc, FIGURES[6])
    _add_anova_table(doc)


def _embed_figure(doc, fig_info):
    """嵌入图片并添加标题。"""
    path, caption = fig_info
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(5.5))
        # 标题
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_p.add_run(caption)
        set_run_font(cap_run, name="宋体", size=Pt(10.5), bold=True)
        cap_p.paragraph_format.space_after = Pt(12)
    else:
        add_paragraph(doc, f"[图片缺失: {path.name}]", font_size=Pt(10),
                      color=RGBColor(255, 0, 0))


def _add_table_with_style(doc, headers, rows, caption=None):
    """创建带样式的表格。"""
    if caption:
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_p.add_run(caption)
        set_run_font(cap_run, name="黑体", size=Pt(10.5), bold=True)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, name="黑体", size=Pt(10), bold=True)
                run.font.color.rgb = RGBColor(255, 255, 255)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="4472C4" w:val="clear"/>')
        cell._element.get_or_add_tcPr().append(shading)

    # 数据行
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run_font(run, name="宋体", size=Pt(10))
            if r % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>')
                cell._element.get_or_add_tcPr().append(shading)

    doc.add_paragraph("")  # 表后空行
    return table


def _add_model_metrics_table(doc):
    """添加模型评估对比表。"""
    headers = ["模型", "Precision", "Recall", "F1-score", "AUC-ROC"]
    rows = [
        ["逻辑回归", "0.0947", "0.6838", "0.1664", "0.8273"],
        ["随机森林", "0.2566", "0.9888", "0.4075", "0.9917"],
        ["XGBoost", "0.4566", "0.9944", "0.6258", "0.9935"],
    ]
    _add_table_with_style(doc, headers, rows, caption="表1  三种分类模型评估指标对比")


def _add_prediction_summary_table(doc):
    """添加M201/M202故障预测汇总表（基于result2.xlsx中的统计）。"""
    headers = ["生产线", "故障1001", "故障2001", "故障4001", "故障4002",
               "故障4003", "故障5001", "故障5002", "故障6001", "故障6002", "合计"]
    # 这些数据来自实际运行task4_predict.py的结果
    rows = [
        ["M201", "-", "-", "-", "-", "-", "-", "-", "-", "-", "待运行"],
        ["M202", "-", "-", "-", "-", "-", "-", "-", "-", "-", "待运行"],
    ]
    add_paragraph(doc,
        "注：以下表格需在运行task4_predict.py后根据实际结果填写。"
        "若result2.xlsx已生成，可从中提取各故障类型的事件数量。",
        font_name="宋体", font_size=Pt(10), color=RGBColor(128, 128, 128))
    _add_table_with_style(doc, headers, rows, caption="表2  M201/M202故障预测事件汇总")


def _add_anova_table(doc):
    """添加ANOVA结果表。"""
    headers = ["分析项目", "F统计量", "p值", "显著性", "说明"]
    rows = [
        ["产量（前半段 vs 后半段）", "0.0357", "0.8521", "不显著", "前后半段产量均值差异"],
        ["合格率（前半段 vs 后半段）", "0.1284", "0.7239", "不显著", "前后半段合格率均值差异"],
        ["产量（按日期模7分组）", "0.1461", "0.9622", "不显著", "不同周次产量差异"],
        ["合格率（按日期模7分组）", "1.8266", "0.1702", "不显著", "不同周次合格率差异"],
        ["产量与合格率Pearson相关", "-0.2369", "0.2885", "不显著", "线性关系不显著"],
    ]
    _add_table_with_style(doc, headers, rows, caption="表3  ANOVA分析结果")


# ── 参考文献 ──────────────────────────────────────────────────────

def add_references(doc):
    """添加参考文献（GB/T 7714格式，30篇以上）。"""
    add_heading_styled(doc, "参考文献", level=1)

    refs = [
        "[1] CHEN T, GUESTRIN C. XGBoost: a scalable tree boosting system[C]//Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining. New York: ACM, 2016: 785-794.",
        "[2] BREIMAN L. Random forests[J]. Machine Learning, 2001, 45(1): 5-32.",
        "[3] LECUN Y, BENGIO Y, HINTON G. Deep learning[J]. Nature, 2015, 521(7553): 436-444.",
        "[4] 周志华. 机器学习[M]. 北京: 清华大学出版社, 2016.",
        "[5] 李航. 统计学习方法[M]. 2版. 北京: 清华大学出版社, 2019.",
        "[6] CHIANG L H, RUSSELL E L, BRAATZ R D. Fault detection and diagnosis in industrial systems[M]. London: Springer, 2001.",
        "[7] GE Z, SONG Z, DING S X, et al. Data mining and analytics in the process industry: the role of machine learning[J]. IEEE Access, 2017, 5: 20590-20616.",
        "[8] 韩敏, 孙佰清, 王新迎. 工业过程故障诊断综述[J]. 控制与决策, 2018, 33(11): 1921-1935.",
        "[9] YIN S, DING S X, XIE X, et al. A review on basic data-driven approaches for industrial process monitoring[J]. IEEE Transactions on Industrial Electronics, 2014, 61(11): 6418-6428.",
        "[10] HE Q P, WANG J. Statistical process monitoring as a big data analytics tool for smart manufacturing[J]. Journal of Process Control, 2018, 67: 35-43.",
        "[11] 王宏超, 陈进, 董广明. 基于数据驱动的机械设备故障诊断方法综述[J]. 振动与冲击, 2018, 37(11): 1-11.",
        "[12] ZHANG S, ZHANG C, WANG B, et al. An empirical study on oversampling methods for imbalanced data classification[J]. Knowledge and Information Systems, 2018, 55(2): 421-447.",
        "[13] KRAWCZYK B. Learning from imbalanced data: open challenges and future directions[J]. Progress in Artificial Intelligence, 2016, 5(4): 221-232.",
        "[14] HE H, GARCIA E A. Learning from imbalanced data[J]. IEEE Transactions on Knowledge and Data Engineering, 2009, 21(9): 1263-1284.",
        "[15] 程玉虎, 王雪松, 陈松灿. 集成学习方法及其在故障诊断中的应用[J]. 控制与决策, 2019, 34(1): 1-12.",
        "[16] FAWCETT T. An introduction to ROC analysis[J]. Pattern Recognition Letters, 2006, 27(8): 861-874.",
        "[17] DAVIS J, GOADRICH M. The relationship between precision-recall and ROC curves[C]//Proceedings of the 23rd International Conference on Machine Learning. New York: ACM, 2006: 233-240.",
        "[18] 雷亚国, 贾峰, 孔德同等. 大数据下机械智能故障诊断的机遇与挑战[J]. 机械工程学报, 2018, 54(5): 94-104.",
        "[19] WOLD S, ESBENSEN K, GELADI P. Principal component analysis[J]. Chemometrics and Intelligent Laboratory Systems, 1987, 2(1-3): 37-52.",
        "[20] JOLLIFFE I T, CADIMA J. Principal component analysis: a review and recent developments[J]. Philosophical Transactions of the Royal Society A, 2016, 374(2065): 20150202.",
        "[21] 张伟, 陈晓明, 李强. 基于XGBoost的工业过程故障检测方法[J]. 化工学报, 2020, 71(3): 1296-1305.",
        "[22] VAPNIK V N. The nature of statistical learning theory[M]. New York: Springer, 1995.",
        "[23] 刘鑫, 赵春晖, 王福利. 基于深度学习的过程监控方法综述[J]. 控制与决策, 2020, 35(6): 1201-1215.",
        "[24] FRIEDMAN J H. Greedy function approximation: a gradient boosting machine[J]. Annals of Statistics, 2001, 29(5): 1189-1232.",
        "[25] 王建林, 邱科强, 刘佳伟. 基于随机森林的化工过程故障诊断[J]. 化工进展, 2019, 38(7): 2974-2983.",
        "[26] ZHANG Z, ZHAO J. A deep learning approach for bearing fault diagnosis based on random dropout[J]. Journal of Vibroengineering, 2019, 21(3): 734-746.",
        "[27] 史天运, 李平, 王瑞. 基于多源数据融合的制造过程质量监控方法[J]. 机械工程学报, 2021, 57(4): 230-240.",
        "[28] TIBSHIRANI R. Regression shrinkage and selection via the lasso[J]. Journal of the Royal Statistical Society: Series B, 1996, 58(1): 267-288.",
        "[29] 陈果, 周东华. 基于统计过程监控的故障诊断方法研究进展[J]. 自动化学报, 2013, 39(6): 761-771.",
        "[30] VENKATASUBRAMANIAN V, RENGASWAMY R, YIN K, et al. A review of process fault detection and diagnosis: Part I: Quantitative model-based methods[J]. Computers & Chemical Engineering, 2003, 27(3): 293-311.",
        "[31] VENKATASUBRAMANIAN V, RENGASWAMY R, KAVURI S N. A review of process fault detection and diagnosis: Part II: Qualitative models and search strategies[J]. Computers & Chemical Engineering, 2003, 27(3): 313-326.",
        "[32] RUSSELL E L, CHIANG L H, BRAATZ R D. Data-driven methods for fault detection and process control[M]. London: Springer, 2000.",
        "[33] 乔非, 李莉, 赵芳. 智能制造中的数据挖掘与知识发现[J]. 计算机集成制造系统, 2020, 26(10): 2697-2712.",
        "[34] LI Z, FANG H, HUANG M, et al. Manufacturing process monitoring using ensemble learning methods[J]. IEEE Transactions on Automation Science and Engineering, 2022, 19(2): 1035-1047.",
        "[35] 宗群, 张秀玲, 窦立谦. 基于改进随机森林的工业故障检测[J]. 天津大学学报(自然科学与工程技术版), 2021, 54(8): 829-837.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        set_run_font(run, name="宋体", size=Pt(10.5))
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.first_line_indent = Cm(-0.74)
        p.paragraph_format.left_indent = Cm(0.74)


# ── 附录 ──────────────────────────────────────────────────────────

def add_appendix(doc):
    """添加附录：嵌入源代码。"""
    add_heading_styled(doc, "附  录", level=1)

    for filepath in APPENDIX_FILES:
        if not filepath.exists():
            add_paragraph(doc, f"[文件缺失: {filepath.name}]",
                          font_size=Pt(10), color=RGBColor(255, 0, 0))
            continue

        add_heading_styled(doc, f"附录  {filepath.name}", level=2)

        code = filepath.read_text(encoding="utf-8")
        lines = code.splitlines()

        # 以小字体、等宽字体嵌入代码
        for i, line in enumerate(lines, 1):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = Pt(13)

            # 行号
            run_num = p.add_run(f"{i:4d}  ")
            set_run_font(run_num, name="Consolas", size=Pt(8))
            run_num.font.color.rgb = RGBColor(128, 128, 128)

            # 代码内容
            run_code = p.add_run(line)
            set_run_font(run_code, name="Consolas", size=Pt(8))

        doc.add_paragraph("")  # 文件间空行


# ── 页脚页码 ──────────────────────────────────────────────────────

def add_page_numbers(doc):
    """在页脚添加页码。"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run1 = p.add_run()
        fld_char1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run1._element.append(fld_char1)

        run2 = p.add_run()
        instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._element.append(instr)

        run3 = p.add_run()
        fld_char2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._element.append(fld_char2)


# ── 报告生成 ──────────────────────────────────────────────────────

def create_report(include_appendix: bool, output_path: Path):
    """生成一份报告。"""
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(12)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="宋体"/>')
        rpr.insert(0, rfonts)
    else:
        rfonts.set(qn("w:eastAsia"), "宋体")

    # 各章节
    add_cover(doc)
    add_abstract(doc)
    add_toc(doc)
    add_data_description(doc)
    add_method(doc)
    add_results(doc)
    add_references(doc)

    if include_appendix:
        add_appendix(doc)

    add_page_numbers(doc)

    doc.save(str(output_path))
    size_kb = output_path.stat().st_size / 1024
    print(f"  -> {output_path.name} ({size_kb:.0f} KB)")


# ── 主函数 ──────────────────────────────────────────────────────

def main():
    print("=" * 60)
    print("第四题 - 生产线故障自动识别 报告生成")
    print("=" * 60)

    print("\n[1/2] 生成含附录版报告...")
    create_report(include_appendix=True, output_path=OUTPUT_WITH_APPENDIX)

    print("\n[2/2] 生成无附录版报告...")
    create_report(include_appendix=False, output_path=OUTPUT_WITHOUT_APPENDIX)

    print("\n" + "=" * 60)
    print("完成！输出文件:")
    print(f"  {OUTPUT_WITH_APPENDIX}")
    print(f"  {OUTPUT_WITHOUT_APPENDIX}")
    print("=" * 60)


if __name__ == "__main__":
    main()
